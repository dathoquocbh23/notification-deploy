#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Sinh file Giai_thich_code_Flutter.docx — giải thích từng file lib theo thứ tự nên tạo."""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

for s in (sys.stdout, sys.stderr):
    try: s.reconfigure(encoding="utf-8")
    except Exception: pass

ROOT = Path(__file__).resolve().parent.parent
APP = ROOT / "app"
NAVY = RGBColor(0x1D, 0x35, 0x57)

doc = Document()

def h1(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
def h2(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
def h3(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.color.rgb = NAVY
def para(t):
    doc.add_paragraph(t)
def bullet(t):
    doc.add_paragraph(t, style="List Bullet") if "List Bullet" in [s.name for s in doc.styles] else doc.add_paragraph("• " + t)
def code_file(rel):
    text = (APP / rel).read_text(encoding="utf-8")
    for line in text.splitlines():
        p = doc.add_paragraph(); r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
def code_block(text):
    for line in text.splitlines():
        p = doc.add_paragraph(); r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5)

# ===== Tiêu đề + tổng quan =====
h1("Giải thích code Flutter — App Ví Demo (phía nhận thông báo)")
para("Tài liệu này giải thích từng file logic trong thư mục app/lib và THỨ TỰ NÊN TẠO. "
     "Nguyên tắc: app KHÔNG chứa nghiệp vụ — nó chỉ là 'tai nghe' làm 3 việc: lấy token, lắng nghe push, "
     "xử lý theo trạng thái app. Mọi điều phối (gửi cho ai, ghi DB) nằm ở backend.")

h2("Thứ tự nên tạo (từ nền tảng lên UI)")
bullet("0. pubspec.yaml — khai báo thư viện (firebase_core, firebase_messaging, flutter_local_notifications, http).")
bullet("1. services/api_service.dart — cầu nối HTTP tới backend (không phụ thuộc file nào → tạo trước).")
bullet("2. services/local_notification_service.dart — hiện banner khi app đang mở (độc lập).")
bullet("3. services/push_service.dart — nhận FCM (DÙNG api_service + local_notification_service → tạo sau 1 & 2).")
bullet("4. main.dart — điểm vào: khởi tạo 3 service rồi chạy app, vào LoginPage.")
bullet("5. pages/login_page.dart — đăng nhập = gắn token vào user.")
bullet("6. pages/home_page.dart — màn chính: số dư, nút chuyển tiền, logout.")
bullet("7. pages/transfer_page.dart — nhập người nhận + số tiền (B1 chuyển tiền).")
bullet("8. pages/otp_page.dart — nhập OTP, verify (B2).")
bullet("9. pages/history_page.dart — lịch sử thông báo đọc từ DB.")
para("Lý do thứ tự này: file ở dưới KHÔNG import file ở trên, nên tạo từ nền tảng (service) lên UI (pages) "
     "thì lúc nào cũng compile được, không gặp lỗi 'chưa có class'.")

# ===== 0. pubspec =====
h2("0. pubspec.yaml — khai báo thư viện")
para("Tạo file này (hoặc thêm dependencies) đầu tiên rồi chạy 'flutter pub get'. 4 gói cốt lõi:")
bullet("firebase_core: khởi tạo Firebase (bắt buộc trước mọi dịch vụ Firebase).")
bullet("firebase_messaging: nhận push FCM, lấy token, subscribe topic.")
bullet("flutter_local_notifications: tự hiện banner khi app đang mở (FCM không tự hiện ở foreground).")
bullet("http: gọi REST API tới backend.")
code_block("""dependencies:
  firebase_core: ^4.1.1
  firebase_messaging: ^16.0.2
  flutter_local_notifications: ^21.0.0
  http: ^1.2.0""")

# ===== 1. api_service =====
h2("1. services/api_service.dart — ApiService (cầu nối backend)")
h3("Vai trò")
para("Gói mọi lời gọi HTTP tới backend vào một nơi. Singleton (một thể hiện duy nhất) để mọi màn hình dùng chung. "
     "Không chứa logic UI, không phụ thuộc file nào khác → đây là class tạo TRƯỚC TIÊN.")
h3("Code")
code_file("lib/services/api_service.dart")
h3("Giải thích logic")
bullet("ApiService._() + static instance: pattern Singleton — constructor private, chỉ có 1 thể hiện 'instance' dùng toàn app.")
bullet("baseUrl: địa chỉ backend. Hiện trỏ domain Railway (HTTPS) nên đi qua internet, không lệ thuộc LAN.")
bullet("currentEmail: lưu email user đang đăng nhập (auth giả lập, không JWT).")
bullet("_post(path, body): hàm dùng chung — POST JSON, đọc kết quả bằng utf8.decode(res.bodyBytes) để không vỡ tiếng Việt.")
bullet("login(): gửi email + displayName + FCM token + deviceInfo lên /auth/login (backend INSERT vào user_tokens), rồi nhớ currentEmail.")
bullet("logout(): gọi /auth/logout để backend XOÁ token thiết bị này (chốt bảo mật), xoá currentEmail.")
bullet("refreshFcmToken(old,new): khi FCM tự đổi token, báo backend UPDATE.")
bullet("createTransfer/verifyTransfer: gọi 2 bước chuyển tiền (/demo/transfers và /verify).")
bullet("notifications()/markRead(): đọc lịch sử thông báo từ DB và đánh dấu đã đọc.")

# ===== 2. local notification =====
h2("2. services/local_notification_service.dart — LocalNotificationService")
h3("Vai trò")
para("Khi app ĐANG MỞ (foreground), FCM KHÔNG tự hiện banner. Service này hiện banner bằng local notification. "
     "Độc lập, không phụ thuộc file app khác → tạo sớm.")
h3("Code")
code_file("lib/services/local_notification_service.dart")
h3("Giải thích logic")
bullet("Singleton giống ApiService.")
bullet("_channel (AndroidNotificationChannel 'banking_channel', Importance.max): Android 8+ bắt buộc notification nằm trên một 'channel'; Importance.max để banner NHẢY heads-up.")
bullet("init(): khởi tạo plugin (icon app), tạo channel, xin quyền thông báo (Android 13+). Lưu ý API v21 dùng named params: initialize(settings:), show(id:..., notificationDetails:...).")
bullet("show(title, body): hiện 1 banner — id lấy theo thời gian để mỗi cái là 1 notification riêng; gắn channel Importance.max + priority.high.")

# ===== 3. push_service =====
h2("3. services/push_service.dart — PushService (trái tim phía nhận)")
h3("Vai trò")
para("Nơi nhận FCM. Làm đúng 3 việc: (1) xin quyền + lấy token, (2) xử lý theo 3 trạng thái app, "
     "(3) báo server khi token đổi. DÙNG ApiService + LocalNotificationService nên tạo SAU hai cái đó.")
h3("Code")
code_file("lib/services/push_service.dart")
h3("Giải thích logic")
bullet("firebaseBackgroundHandler (top-level + @pragma('vm:entry-point')): handler chạy khi app ở nền/đã tắt, trong isolate riêng. BẮT BUỘC để ngoài class + có @pragma để không bị release build tree-shake mất. Không cần tự hiện gì — HỆ ĐIỀU HÀNH tự hiện notification message.")
bullet("init() bước (1): requestPermission() xin quyền; getToken() = 'địa chỉ' của máy này (~142 ký tự). Token này gửi lên server lúc login.")
bullet("subscribeToTopic('all'): đăng ký 'đài' broadcast — backend bắn tới topic 'all' là máy này nhận (luồng khuyến mãi 1-nhiều).")
bullet("(3) onTokenRefresh.listen: FCM có thể tự đổi token; nghe sự kiện này -> gọi ApiService.refreshFcmToken để server UPDATE.")
bullet("(2a) onMessage.listen — FOREGROUND: app đang mở, FCM không tự hiện -> gọi LocalNotificationService.show() bắc cầu.")
bullet("(2b) onMessageOpenedApp — BACKGROUND: user chạm noti khi app chạy nền -> gọi onOpenHistory() để mở màn lịch sử.")
bullet("(2c) getInitialMessage — TERMINATED: app vừa được MỞ do user chạm noti lúc app đã tắt -> điều hướng vào lịch sử.")
bullet("onBackgroundMessage(firebaseBackgroundHandler): đăng ký handler nền ở cuối.")
bullet("onOpenHistory: callback để UI (HomePage) gán hành vi điều hướng — service không biết gì về widget.")

# ===== 4. main =====
h2("4. main.dart — điểm vào + vỏ app")
h3("Code")
code_file("lib/main.dart")
h3("Giải thích logic")
bullet("WidgetsFlutterBinding.ensureInitialized(): bắt buộc trước khi gọi code native (Firebase, plugin) trong main async.")
bullet("Mỗi init bọc try-catch: nếu 1 service hỏng (vd getToken trả SERVICE_NOT_AVAILABLE khi mạng chặn Google) thì CHỈ log, KHÔNG để exception làm main() dừng trước runApp() -> tránh màn hình trắng.")
bullet("Thứ tự init: Firebase.initializeApp() trước (mọi dịch vụ Firebase cần) -> LocalNotificationService -> PushService.")
bullet("runApp(NotifyDemoApp): luôn được gọi -> UI luôn vẽ. NotifyDemoApp là MaterialApp, home = LoginPage.")

# ===== 5..9 pages =====
PAGES = [
 ("5. pages/login_page.dart — LoginPage", "lib/pages/login_page.dart",
  "Đăng nhập = thời điểm GẮN TOKEN VÀO USER (SD-1).",
  ["StatefulWidget vì có ô nhập + trạng thái _loading.",
   "_login(): lấy token từ PushService; NẾU token null thì return (chưa có token thì không đăng ký được) -> lưu ý: máy chưa lấy được FCM token thì bấm sẽ 'không có gì xảy ra'.",
   "Gọi ApiService.login(... fcmToken: token ...) -> backend INSERT user_tokens; rồi pushReplacement sang HomePage.",
   "mounted check trước khi dùng context (tránh lỗi khi widget đã bị huỷ).",
   "build(): 2 TextField (email, tên) + nút Đăng nhập (disable khi đang loading)."]),
 ("6. pages/home_page.dart — HomePage", "lib/pages/home_page.dart",
  "Màn chính sau đăng nhập: số dư giả lập, nút chuyển tiền, chuông lịch sử, logout.",
  ["initState() gán PushService.instance.onOpenHistory = mở HistoryPage -> khi user chạm noti (background/terminated), service gọi callback này để điều hướng.",
   "_logout(): gọi ApiService.logout(token) để backend XOÁ token thiết bị (không xoá thì máy cũ vẫn nhận noti của tài khoản cũ) rồi về LoginPage.",
   "build(): AppBar hiện email + nút chuông (HistoryPage) + nút logout; Card số dư; nút Chuyển tiền -> TransferPage; hiện token để demo."]),
 ("7. pages/transfer_page.dart — TransferPage", "lib/pages/transfer_page.dart",
  "Bước 1 chuyển tiền: nhập email người nhận + số tiền -> backend sinh OTP gửi email.",
  ["_submit(): gọi ApiService.createTransfer(toEmail, amount); nếu trả về transferId -> sang OtpPage kèm transferId.",
   "Mặc định số tiền 100000 cho nhanh khi demo.",
   "build(): 2 TextField (email người nhận, số tiền) + nút Tiếp tục."]),
 ("8. pages/otp_page.dart — OtpPage", "lib/pages/otp_page.dart",
  "Bước 2: nhập OTP (lấy từ email) -> verify -> backend fan-out push + email biên lai.",
  ["Nhận transferId qua constructor (từ TransferPage).",
   "_verify(): gọi ApiService.verifyTransfer(transferId, otp); nếu status SUCCESS -> pop về và hiện SnackBar 'chờ thông báo đẩy về'; nếu sai -> hiện errorText.",
   "build(): ô nhập 6 số (maxLength 6, căn giữa, giãn chữ) + nút Xác nhận."]),
 ("9. pages/history_page.dart — HistoryPage", "lib/pages/history_page.dart",
  "Lịch sử thông báo đọc từ DB (nguồn chân lý), KHÔNG từ push (SD-3).",
  ["initState() gọi _load() -> ApiService.notifications() đọc danh sách từ backend (đọc từ bảng notifications).",
   "Vì đọc từ DB nên đủ cả thông báo mà push bị rớt (token chết, mất mạng, tắt quyền).",
   "_icon(type): chọn icon theo loại TRANSFER_IN/OUT/PROMO.",
   "build(): ListView các ListTile; chưa đọc (readAt==null) -> in đậm + chấm đỏ; chạm -> markRead rồi load lại; kéo xuống để refresh."]),
]
for title, rel, role, bullets in PAGES:
    h2(title)
    h3("Vai trò"); para(role)
    h3("Code"); code_file(rel)
    h3("Giải thích logic")
    for b in bullets: bullet(b)

# ===== Sơ đồ luồng =====
h2("Sơ đồ luồng tổng (ai gọi ai)")
code_block("""main()  --init-->  LocalNotificationService, PushService
                         |               |
   LoginPage --login--> ApiService --HTTP--> Backend (INSERT user_tokens)
   TransferPage/OtpPage --HTTP--> Backend (transfer, verify) --> fan-out FCM
   Backend --FCM push--> PushService:
        foreground -> LocalNotificationService.show() (banner)
        background/terminated -> OS hiện banner -> onOpenHistory -> HistoryPage
   HistoryPage --HTTP--> Backend (đọc bảng notifications = nguồn chân lý)""")

out = ROOT / "Giai_thich_code_Flutter.docx"
doc.save(str(out))
print("Đã tạo:", out)
print("Số paragraph:", len(doc.paragraphs))
