#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Cỗ máy báo cáo — append một BƯỚC vào Bao_cao_thuc_hien.docx.

Dùng:
    python scripts/report.py --buoc "BƯỚC 1 — Tên bước" --file noidung.md

Hành vi:
  * heading lớn = tên bước (style "heading 1", fallback bold 17pt navy RGB(29,53,87));
  * đọc markdown:
      - dòng "## ..."  -> style "heading 3" (fallback bold navy)
      - dòng "- ..."   -> bullet
      - dòng "BUG:..." -> paragraph đỏ RGB(192,57,43) đậm
      - còn lại        -> paragraph thường
  * nếu có report_assets/<buoc-slug>/ -> nhúng mọi .png/.jpg (căn giữa, rộng 15cm);
  * lưu đè file docx.
"""
import argparse
import re
import sys
import unicodedata
from pathlib import Path

# Console Windows mặc định cp1252 -> ép UTF-8 để in tiếng Việt không vỡ.
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8")
    except Exception:
        pass

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Mặc định nằm cùng repo root (thư mục cha của scripts/)
ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "Bao_cao_thuc_hien.docx"
ASSETS_ROOT = ROOT / "report_assets"

NAVY = RGBColor(0x1D, 0x35, 0x57)   # RGB(29,53,87)
BUG_RED = RGBColor(0xC0, 0x39, 0x2B)  # RGB(192,57,43)
IMAGE_EXTS = (".png", ".jpg", ".jpeg")


def slugify(text):
    """'BƯỚC 0 — Dọn bảo mật' -> 'buoc-0-don-bao-mat' (bỏ dấu, ascii, gạch nối)."""
    text = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in text if not unicodedata.combining(c))
    text = text.replace("đ", "d").replace("Đ", "d")
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return text.strip("-")


def find_style(doc, *names):
    """Tìm style theo tên (không phân biệt hoa/thường). Trả về object hoặc None."""
    wanted = {n.lower() for n in names}
    for style in doc.styles:
        try:
            if style.name and style.name.lower() in wanted:
                return style
        except Exception:
            continue
    return None


def add_heading1(doc, text):
    style = find_style(doc, "heading 1", "Heading 1")
    if style is not None:
        doc.add_paragraph(text, style=style)
    else:  # fallback: bold 17pt navy
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(17)
        run.font.color.rgb = NAVY


def add_heading3(doc, text):
    style = find_style(doc, "heading 3", "Heading 3")
    if style is not None:
        doc.add_paragraph(text, style=style)
    else:  # fallback: bold navy
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = NAVY


def add_bullet(doc, text):
    style = find_style(doc, "list bullet", "List Bullet")
    if style is not None:
        doc.add_paragraph(text, style=style)
    else:  # fallback: chấm đầu dòng thủ công
        doc.add_paragraph("• " + text)


def add_bug(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = BUG_RED


def asset_dirs(buoc):
    """Thư mục ảnh ứng viên: slug đầy đủ + dạng ngắn 'buoc<N>' nếu tên bước có số."""
    slug = slugify(buoc)
    dirs = [ASSETS_ROOT / slug]
    m = re.match(r"buoc-?(\d+)", slug)
    if m:
        short = ASSETS_ROOT / f"buoc{m.group(1)}"
        if short not in dirs:
            dirs.append(short)
    return dirs


def add_images(doc, buoc):
    imgs = []
    for folder in asset_dirs(buoc):
        if folder.is_dir():
            imgs += sorted(
                p for p in folder.iterdir()
                if p.suffix.lower() in IMAGE_EXTS
            )
    for img in imgs:
        doc.add_picture(str(img), width=Cm(15))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return len(imgs)


def add_code(doc, text):
    """Một dòng lệnh/code/log — font monospace nhỏ để chứa dòng dài, phân biệt văn bản thường."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def render_markdown(doc, md_path):
    text = Path(md_path).read_text(encoding="utf-8")
    in_code = False
    for raw in text.splitlines():
        line = raw.rstrip()
        # Khối code ```...``` -> render monospace, giữ thụt đầu dòng
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            add_code(doc, line)
            continue
        if not line.strip():
            continue
        if line.startswith("## "):
            add_heading3(doc, line[3:].strip())
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        elif line.lstrip().startswith("BUG:"):
            add_bug(doc, line.strip())
        else:
            doc.add_paragraph(line)


def append_section(buoc, md_file, docx_path=DOCX):
    docx_path = Path(docx_path)
    doc = Document(str(docx_path)) if docx_path.exists() else Document()
    add_heading1(doc, buoc)
    render_markdown(doc, md_file)
    n_imgs = add_images(doc, buoc)
    doc.save(str(docx_path))
    return n_imgs


def main():
    ap = argparse.ArgumentParser(description="Append một BƯỚC vào báo cáo .docx")
    ap.add_argument("--buoc", required=True, help='Tên bước, vd "BƯỚC 1 — ..."')
    ap.add_argument("--file", required=True, help="Đường dẫn file markdown nội dung")
    ap.add_argument("--docx", default=str(DOCX), help="File docx đích (mặc định: Bao_cao_thuc_hien.docx)")
    args = ap.parse_args()

    if not Path(args.file).exists():
        sys.exit(f"Không tìm thấy file nội dung: {args.file}")

    n = append_section(args.buoc, args.file, args.docx)
    cands = " hoặc ".join(str(d) for d in asset_dirs(args.buoc))
    print(f"✔ Đã append '{args.buoc}' vào {args.docx}")
    print(f"  - nhúng {n} ảnh" if n else f"  - không có ảnh (đặt vào: {cands})")


if __name__ == "__main__":
    main()
